"""
Document Ingestion Pipeline.
Recursively loads files from data/documents, processes them, and chunks them into retrieval-ready format.

Supports: PDF (.pdf), Text (.txt), Word Documents (.docx)
Outputs: List of dicts with id, text, and metadata for each chunk.
"""
import os
import re
import logging
from pathlib import Path
from typing import List, Dict, Any, Union

# Import document processing libraries
try:
    from PyPDF2 import PdfReader
except ImportError:
    PdfReader = None

try:
    from docx import Document as DocxDocument
except ImportError:
    DocxDocument = None

try:
    from langchain.text_splitter import RecursiveCharacterTextSplitter
    LANGCHAIN_AVAILABLE = True
except ImportError:
    LANGCHAIN_AVAILABLE = False

# Configure logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger(__name__)


class IngestionError(Exception):
    """Custom exception for ingestion errors."""
    pass


class DocumentIngestor:
    """
    Handles recursive document loading, text extraction, cleaning, and chunking.
    """
    
    def __init__(
        self,
        documents_dir: str = "data/documents",
        chunk_size_min: int = 500,
        chunk_size_max: int = 800,
        overlap_percent: int = 10
    ):
        """
        Initialize the document ingestor.
        
        Args:
            documents_dir: Directory to recursively load documents from
            chunk_size_min: Minimum characters per chunk
            chunk_size_max: Maximum characters per chunk
            overlap_percent: Percentage overlap between chunks
        """
        self.documents_dir = Path(documents_dir)
        self.chunk_size_min = chunk_size_min
        self.chunk_size_max = chunk_size_max
        self.overlap_percent = overlap_percent
        self.overlap_size = int(chunk_size_max * overlap_percent / 100)
        
        # Supported file extensions
        self.supported_extensions = {'.pdf', '.txt', '.docx'}
        
        logger.info(f"DocumentIngestor initialized: dir={documents_dir}, chunk_size={chunk_size_min}-{chunk_size_max}, overlap={overlap_percent}%")
    
    def _clean_text(self, text: str) -> str:
        """
        Clean text by normalizing whitespace and removing null characters.
        
        Args:
            text: Raw text to clean
            
        Returns:
            str: Cleaned text
        """
        # Remove null characters
        text = text.replace('\x00', '')
        text = text.replace('\u0000', '')
        
        # Normalize whitespace: replace multiple spaces/tabs/newlines with single space
        text = re.sub(r'[ \t]+', ' ', text)  # Multiple spaces/tabs -> single space
        text = re.sub(r'\n\s*\n', '\n\n', text)  # Multiple newlines -> double newline (paragraph break)
        text = re.sub(r'\n', ' ', text)  # Single newlines -> space
        text = re.sub(r'\s+', ' ', text)  # Final cleanup: any remaining multiple spaces
        
        # Strip leading/trailing whitespace
        text = text.strip()
        
        return text
    
    def _load_pdf(self, file_path: Path) -> str:
        """
        Extract text from PDF file using PyPDF2.
        
        Args:
            file_path: Path to PDF file
            
        Returns:
            str: Extracted text
        """
        if PdfReader is None:
            raise IngestionError("PyPDF2 not installed. Install with: pip install PyPDF2")
        
        try:
            reader = PdfReader(str(file_path))
            text_parts = []
            
            for page_num, page in enumerate(reader.pages, 1):
                text = page.extract_text()
                if text.strip():
                    text_parts.append(text)
                else:
                    logger.warning(f"No text on page {page_num} of {file_path.name}")
            
            full_text = "\n\n".join(text_parts)
            
            if not full_text.strip():
                raise IngestionError(f"No text extracted from PDF: {file_path.name}")
            
            logger.info(f"Extracted text from PDF: {file_path.name} ({len(reader.pages)} pages)")
            return full_text
            
        except Exception as e:
            raise IngestionError(f"Error loading PDF {file_path.name}: {str(e)}")
    
    def _load_txt(self, file_path: Path) -> str:
        """
        Extract text from TXT file using plain open().
        
        Args:
            file_path: Path to TXT file
            
        Returns:
            str: Extracted text
        """
        try:
            # Try multiple encodings
            encodings = ['utf-8', 'utf-8-sig', 'latin-1', 'cp1252']
            
            for encoding in encodings:
                try:
                    with open(file_path, 'r', encoding=encoding) as f:
                        text = f.read()
                    logger.info(f"Extracted text from TXT: {file_path.name} (encoding: {encoding})")
                    return text
                except UnicodeDecodeError:
                    continue
            
            raise IngestionError(f"Could not decode TXT file {file_path.name} with any supported encoding")
            
        except Exception as e:
            raise IngestionError(f"Error loading TXT {file_path.name}: {str(e)}")
    
    def _load_docx(self, file_path: Path) -> str:
        """
        Extract text from DOCX file using python-docx.
        
        Args:
            file_path: Path to DOCX file
            
        Returns:
            str: Extracted text
        """
        if DocxDocument is None:
            raise IngestionError("python-docx not installed. Install with: pip install python-docx")
        
        try:
            doc = DocxDocument(str(file_path))
            
            # Extract paragraphs
            paragraphs = [para.text for para in doc.paragraphs if para.text.strip()]
            
            # Extract tables
            table_texts = []
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join([cell.text.strip() for cell in row.cells])
                    if row_text.strip():
                        table_texts.append(row_text)
            
            # Combine all text
            full_text = "\n\n".join(paragraphs + table_texts)
            
            if not full_text.strip():
                raise IngestionError(f"No text extracted from DOCX: {file_path.name}")
            
            logger.info(f"Extracted text from DOCX: {file_path.name}")
            return full_text
            
        except Exception as e:
            raise IngestionError(f"Error loading DOCX {file_path.name}: {str(e)}")
    
    def _find_documents(self) -> List[Path]:
        """
        Recursively find all supported documents in the documents directory.
        
        Returns:
            List[Path]: List of document file paths
        """
        if not self.documents_dir.exists():
            logger.warning(f"Documents directory does not exist: {self.documents_dir}")
            return []
        
        documents = []
        
        for ext in self.supported_extensions:
            # Use rglob for recursive search
            found = list(self.documents_dir.rglob(f"*{ext}"))
            documents.extend(found)
        
        logger.info(f"Found {len(documents)} documents in {self.documents_dir}")
        return sorted(documents)
    
    def _split_into_sentences(self, text: str) -> List[str]:
        """
        Split text into sentences for better chunking.
        
        Args:
            text: Input text
            
        Returns:
            List[str]: List of sentences
        """
        # Pattern to split on sentence boundaries
        sentence_pattern = r'(?<=[.!?])\s+(?=[A-Z])|(?<=[.!?])\n+'
        sentences = re.split(sentence_pattern, text)
        return [s.strip() for s in sentences if s.strip()]
    
    def _chunk_text_custom(self, text: str, source: str) -> List[Dict[str, Any]]:
        """
        Custom text splitter: chunks text into 500-800 chars with 10% overlap.
        Falls back to this if langchain is not available.
        
        Args:
            text: Text to chunk
            source: Source filename
            
        Returns:
            List[Dict]: List of chunk dictionaries
        """
        sentences = self._split_into_sentences(text)
        
        if not sentences:
            # If no sentences, treat whole text as one chunk
            sentences = [text]
        
        chunks = []
        chunk_index = 0
        sentence_idx = 0
        
        while sentence_idx < len(sentences):
            # Build chunk from sentences
            chunk_text = ""
            start_idx = sentence_idx
            
            # Add sentences until minimum size
            while sentence_idx < len(sentences) and len(chunk_text) < self.chunk_size_min:
                if chunk_text:
                    chunk_text += " "
                chunk_text += sentences[sentence_idx]
                sentence_idx += 1
            
            # Continue adding if we haven't hit max
            while sentence_idx < len(sentences) and len(chunk_text) < self.chunk_size_max:
                next_sent = sentences[sentence_idx]
                if len(chunk_text) + len(next_sent) + 1 <= self.chunk_size_max:
                    chunk_text += " " + next_sent
                    sentence_idx += 1
                else:
                    break
            
            # Create chunk dict
            chunks.append({
                "id": f"{source}_{chunk_index}",
                "text": chunk_text.strip(),
                "metadata": {
                    "source": source,
                    "chunk_index": chunk_index
                }
            })
            
            # Calculate overlap for next chunk
            if sentence_idx < len(sentences):
                # Go back by overlap_percent
                overlap_chars = 0
                overlap_start = sentence_idx
                
                while overlap_start > start_idx and overlap_chars < self.overlap_size:
                    overlap_start -= 1
                    overlap_chars += len(sentences[overlap_start])
                
                sentence_idx = overlap_start if overlap_start > start_idx else sentence_idx
            
            chunk_index += 1
        
        return chunks
    
    def _chunk_text_langchain(self, text: str, source: str) -> List[Dict[str, Any]]:
        """
        Use langchain RecursiveCharacterTextSplitter for chunking.
        
        Args:
            text: Text to chunk
            source: Source filename
            
        Returns:
            List[Dict]: List of chunk dictionaries
        """
        splitter = RecursiveCharacterTextSplitter(
            chunk_size=self.chunk_size_max,
            chunk_overlap=self.overlap_size,
            length_function=len,
            separators=["\n\n", "\n", ". ", " ", ""]
        )
        
        text_chunks = splitter.split_text(text)
        
        chunks = []
        for idx, chunk_text in enumerate(text_chunks):
            chunks.append({
                "id": f"{source}_{idx}",
                "text": chunk_text.strip(),
                "metadata": {
                    "source": source,
                    "chunk_index": idx
                }
            })
        
        return chunks
    
    def _chunk_text(self, text: str, source: str) -> List[Dict[str, Any]]:
        """
        Split text into chunks with metadata.
        Uses langchain if available, otherwise custom splitter.
        
        Args:
            text: Text to chunk
            source: Source filename
            
        Returns:
            List[Dict]: List of chunks with format:
                {"id": "<source>_<chunk_index>", "text": "...", "metadata": {...}}
        """
        if LANGCHAIN_AVAILABLE:
            logger.debug(f"Using langchain RecursiveCharacterTextSplitter for {source}")
            return self._chunk_text_langchain(text, source)
        else:
            logger.debug(f"Using custom text splitter for {source}")
            return self._chunk_text_custom(text, source)
    
    def load_and_process_file(self, file_path: Union[str, Path]) -> List[Dict[str, Any]]:
        """
        Load a single file, extract text, clean it, and chunk it.
        
        Args:
            file_path: Path to document file
            
        Returns:
            List[Dict]: List of chunk dictionaries
        """
        file_path = Path(file_path)
        ext = file_path.suffix.lower()
        source = file_path.name
        
        try:
            # Load text based on file type
            if ext == '.pdf':
                raw_text = self._load_pdf(file_path)
            elif ext == '.txt':
                raw_text = self._load_txt(file_path)
            elif ext == '.docx':
                raw_text = self._load_docx(file_path)
            else:
                raise IngestionError(f"Unsupported file type: {ext}")
            
            # Clean text
            cleaned_text = self._clean_text(raw_text)
            
            if not cleaned_text:
                logger.warning(f"No text after cleaning: {source}")
                return []
            
            # Chunk text
            chunks = self._chunk_text(cleaned_text, source)
            
            logger.info(f"Processed {source}: {len(cleaned_text)} chars -> {len(chunks)} chunks")
            return chunks
            
        except Exception as e:
            logger.error(f"Failed to process {source}: {str(e)}")
            return []
    
    def ingest_documents(self) -> List[Dict[str, Any]]:
        """
        Main ingestion method: recursively loads all documents from documents_dir,
        processes them, and returns list of chunks.
        
        Returns:
            List[Dict]: List of all chunks from all documents with format:
                {"id": "<source>_<chunk_index>", "text": "...", "metadata": {"source": "...", "chunk_index": i}}
        """
        # Find all documents
        document_paths = self._find_documents()
        
        if not document_paths:
            logger.warning("No documents found to ingest")
            return []
        
        # Process each document
        all_chunks = []
        
        for doc_path in document_paths:
            chunks = self.load_and_process_file(doc_path)
            all_chunks.extend(chunks)
        
        logger.info(f"Ingestion complete: {len(document_paths)} documents -> {len(all_chunks)} total chunks")
        return all_chunks


def ingest_documents(
    documents_dir: str = "data/documents",
    chunk_size_min: int = 500,
    chunk_size_max: int = 800,
    overlap_percent: int = 10
) -> List[Dict[str, Any]]:
    """
    Convenience function to ingest documents with custom parameters.
    
    Args:
        documents_dir: Directory containing documents to ingest
        chunk_size_min: Minimum chunk size (characters)
        chunk_size_max: Maximum chunk size (characters)
        overlap_percent: Overlap between chunks (percent)
        
    Returns:
        List[Dict]: List of chunk dictionaries
    """
    ingestor = DocumentIngestor(
        documents_dir=documents_dir,
        chunk_size_min=chunk_size_min,
        chunk_size_max=chunk_size_max,
        overlap_percent=overlap_percent
    )
    return ingestor.ingest_documents()


if __name__ == "__main__":
    """
    Demo: Load documents, process them, and print statistics.
    """
    print("=" * 80)
    print("Document Ingestion Pipeline Demo")
    print("=" * 80)
    
    # Configuration
    DOCUMENTS_DIR = "data/documents"
    
    print(f"\nLoading documents from: {DOCUMENTS_DIR}")
    print(f"Chunk size: 500-800 characters with 10% overlap")
    print("-" * 80)
    
    # Create data/documents directory if it doesn't exist
    os.makedirs(DOCUMENTS_DIR, exist_ok=True)
    
    # Check if directory has any files
    doc_dir = Path(DOCUMENTS_DIR)
    existing_files = list(doc_dir.glob("**/*"))
    supported_files = [f for f in existing_files if f.suffix.lower() in {'.pdf', '.txt', '.docx'}]
    
    if not supported_files:
        print(f"\n⚠️  No documents found in {DOCUMENTS_DIR}")
        print("   Creating a sample text file for demonstration...")
        
        # Create a sample document
        sample_file = doc_dir / "sample_document.txt"
        sample_content = """
        This is a sample document for the RAG chatbot ingestion pipeline.
        
        The pipeline supports PDF, TXT, and DOCX files. It recursively loads files from the data/documents
        directory, extracts text, cleans it by normalizing whitespace and removing null characters, and then
        splits the text into chunks of 500-800 characters with 10% overlap.
        
        Each chunk is assigned a unique ID in the format: <source>_<chunk_index>. The chunks also include
        metadata such as the source filename and chunk index.
        
        This chunking strategy ensures efficient retrieval by creating manageable text segments that can be
        embedded and stored in a vector database. The overlap between chunks helps maintain context across
        chunk boundaries, which is important for answering questions that might span multiple chunks.
        
        The ingestion pipeline includes error handling and logging to track the processing of each document.
        It can handle multiple file formats and encodings, making it robust for real-world document processing.
        """
        
        with open(sample_file, 'w', encoding='utf-8') as f:
            f.write(sample_content)
        
        print(f"   ✓ Created: {sample_file}")
    
    # Run ingestion
    try:
        chunks = ingest_documents(documents_dir=DOCUMENTS_DIR)
        
        print(f"\n{'=' * 80}")
        print("INGESTION RESULTS")
        print("=" * 80)
        print(f"Total chunks created: {len(chunks)}")
        
        if chunks:
            # Calculate statistics
            total_chars = sum(len(chunk['text']) for chunk in chunks)
            avg_chunk_size = total_chars / len(chunks) if chunks else 0
            
            sources = set(chunk['metadata']['source'] for chunk in chunks)
            
            print(f"Documents processed: {len(sources)}")
            print(f"Average chunk size: {avg_chunk_size:.1f} characters")
            print(f"Total text: {total_chars:,} characters")
            
            # Print sample chunk
            print(f"\n{'-' * 80}")
            print("SAMPLE CHUNK (first chunk):")
            print("-" * 80)
            sample = chunks[0]
            print(f"ID: {sample['id']}")
            print(f"Source: {sample['metadata']['source']}")
            print(f"Chunk Index: {sample['metadata']['chunk_index']}")
            print(f"Text Length: {len(sample['text'])} characters")
            print(f"\nText Preview:")
            print(sample['text'][:300] + "..." if len(sample['text']) > 300 else sample['text'])
            
            # Print chunk count by source
            print(f"\n{'-' * 80}")
            print("CHUNKS BY SOURCE:")
            print("-" * 80)
            from collections import Counter
            source_counts = Counter(chunk['metadata']['source'] for chunk in chunks)
            for source, count in sorted(source_counts.items()):
                print(f"  {source}: {count} chunks")
        else:
            print("\n⚠️  No chunks were created. Check that documents contain extractable text.")
        
        print(f"\n{'=' * 80}")
        print("Demo complete!")
        print("=" * 80)
        
    except Exception as e:
        logger.error(f"Error during ingestion: {str(e)}", exc_info=True)
        print(f"\n❌ Error: {str(e)}")
