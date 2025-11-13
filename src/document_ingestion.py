"""
Document Ingestion Module.
Handles loading and extracting text from PDF, TXT, and DOCX files.
"""
import os
from pathlib import Path
from typing import Union, List, Tuple
import logging

# PDF processing
try:
    from PyPDF2 import PdfReader
except ImportError:
    PdfReader = None

# DOCX processing
try:
    from docx import Document
except ImportError:
    Document = None

from .config import SUPPORTED_EXTENSIONS, MAX_FILE_SIZE_MB

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)


class DocumentIngestionError(Exception):
    """Custom exception for document ingestion errors."""
    pass


class DocumentIngestor:
    """
    Handles ingestion of various document formats (PDF, TXT, DOCX).
    Extracts text content and metadata from uploaded files.
    """
    
    def __init__(self):
        """Initialize the document ingestor."""
        self.supported_extensions = SUPPORTED_EXTENSIONS
        self.max_file_size_mb = MAX_FILE_SIZE_MB
    
    def validate_file(self, file_path: Union[str, Path]) -> bool:
        """
        Validate that the file exists, has supported extension, and is within size limits.
        
        Args:
            file_path: Path to the file to validate
            
        Returns:
            bool: True if file is valid
            
        Raises:
            DocumentIngestionError: If file is invalid
        """
        file_path = Path(file_path)
        
        # Check if file exists
        if not file_path.exists():
            raise DocumentIngestionError(f"File not found: {file_path}")
        
        # Check file extension
        if file_path.suffix.lower() not in self.supported_extensions:
            raise DocumentIngestionError(
                f"Unsupported file type: {file_path.suffix}. "
                f"Supported types: {', '.join(self.supported_extensions)}"
            )
        
        # Check file size
        file_size_mb = file_path.stat().st_size / (1024 * 1024)
        if file_size_mb > self.max_file_size_mb:
            raise DocumentIngestionError(
                f"File too large: {file_size_mb:.2f}MB. "
                f"Maximum allowed: {self.max_file_size_mb}MB"
            )
        
        return True
    
    def ingest_pdf(self, file_path: Union[str, Path]) -> str:
        """
        Extract text from a PDF file.
        
        Args:
            file_path: Path to the PDF file
            
        Returns:
            str: Extracted text content
            
        Raises:
            DocumentIngestionError: If PDF processing fails
        """
        if PdfReader is None:
            raise DocumentIngestionError("PyPDF2 not installed. Install with: pip install PyPDF2")
        
        try:
            file_path = Path(file_path)
            reader = PdfReader(str(file_path))
            
            text_content = []
            for page_num, page in enumerate(reader.pages, 1):
                text = page.extract_text()
                if text.strip():
                    text_content.append(text)
                else:
                    logger.warning(f"No text extracted from page {page_num} of {file_path.name}")
            
            full_text = "\n\n".join(text_content)
            
            if not full_text.strip():
                raise DocumentIngestionError(f"No text could be extracted from PDF: {file_path.name}")
            
            logger.info(f"Successfully extracted text from PDF: {file_path.name} ({len(reader.pages)} pages)")
            return full_text
            
        except Exception as e:
            raise DocumentIngestionError(f"Error processing PDF {file_path.name}: {str(e)}")
    
    def ingest_txt(self, file_path: Union[str, Path]) -> str:
        """
        Extract text from a TXT file.
        
        Args:
            file_path: Path to the TXT file
            
        Returns:
            str: Extracted text content
            
        Raises:
            DocumentIngestionError: If TXT processing fails
        """
        try:
            file_path = Path(file_path)
            
            # Try different encodings
            encodings = ['utf-8', 'utf-8-sig', 'latin-1', 'cp1252']
            text_content = None
            
            for encoding in encodings:
                try:
                    with open(file_path, 'r', encoding=encoding) as f:
                        text_content = f.read()
                    break  # Success, exit loop
                except UnicodeDecodeError:
                    continue
            
            if text_content is None:
                raise DocumentIngestionError(
                    f"Could not decode text file {file_path.name} with any supported encoding"
                )
            
            if not text_content.strip():
                raise DocumentIngestionError(f"Text file is empty: {file_path.name}")
            
            logger.info(f"Successfully extracted text from TXT: {file_path.name}")
            return text_content
            
        except Exception as e:
            if isinstance(e, DocumentIngestionError):
                raise
            raise DocumentIngestionError(f"Error processing TXT {file_path.name}: {str(e)}")
    
    def ingest_docx(self, file_path: Union[str, Path]) -> str:
        """
        Extract text from a DOCX file.
        
        Args:
            file_path: Path to the DOCX file
            
        Returns:
            str: Extracted text content
            
        Raises:
            DocumentIngestionError: If DOCX processing fails
        """
        if Document is None:
            raise DocumentIngestionError("python-docx not installed. Install with: pip install python-docx")
        
        try:
            file_path = Path(file_path)
            doc = Document(str(file_path))
            
            # Extract text from paragraphs
            paragraphs = [para.text for para in doc.paragraphs if para.text.strip()]
            
            # Extract text from tables
            table_texts = []
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join([cell.text.strip() for cell in row.cells])
                    if row_text.strip():
                        table_texts.append(row_text)
            
            # Combine all text
            all_text = paragraphs + table_texts
            full_text = "\n\n".join(all_text)
            
            if not full_text.strip():
                raise DocumentIngestionError(f"No text could be extracted from DOCX: {file_path.name}")
            
            logger.info(f"Successfully extracted text from DOCX: {file_path.name}")
            return full_text
            
        except Exception as e:
            if isinstance(e, DocumentIngestionError):
                raise
            raise DocumentIngestionError(f"Error processing DOCX {file_path.name}: {str(e)}")
    
    def ingest_document(self, file_path: Union[str, Path]) -> Tuple[str, dict]:
        """
        Main ingestion method that routes to appropriate handler based on file type.
        
        Args:
            file_path: Path to the document file
            
        Returns:
            Tuple[str, dict]: Extracted text content and metadata
            
        Raises:
            DocumentIngestionError: If ingestion fails
        """
        file_path = Path(file_path)
        
        # Validate file
        self.validate_file(file_path)
        
        # Extract text based on file type
        extension = file_path.suffix.lower()
        
        if extension == ".pdf":
            text = self.ingest_pdf(file_path)
        elif extension == ".txt":
            text = self.ingest_txt(file_path)
        elif extension == ".docx":
            text = self.ingest_docx(file_path)
        else:
            raise DocumentIngestionError(f"Unsupported file type: {extension}")
        
        # Create metadata
        metadata = {
            "source": file_path.name,
            "file_type": extension.lstrip('.'),
            "file_size_bytes": file_path.stat().st_size,
            "char_count": len(text)
        }
        
        logger.info(f"Document ingestion complete: {file_path.name} - {len(text)} characters")
        
        return text, metadata


def ingest_document(file_path: Union[str, Path]) -> Tuple[str, dict]:
    """
    Convenience function to ingest a single document.
    
    Args:
        file_path: Path to the document file
        
    Returns:
        Tuple[str, dict]: Extracted text content and metadata
    """
    ingestor = DocumentIngestor()
    return ingestor.ingest_document(file_path)


def ingest_documents(file_paths: List[Union[str, Path]]) -> List[Tuple[str, dict]]:
    """
    Ingest multiple documents.
    
    Args:
        file_paths: List of paths to document files
        
    Returns:
        List[Tuple[str, dict]]: List of (text, metadata) tuples
    """
    ingestor = DocumentIngestor()
    results = []
    
    for file_path in file_paths:
        try:
            text, metadata = ingestor.ingest_document(file_path)
            results.append((text, metadata))
        except DocumentIngestionError as e:
            logger.error(f"Failed to ingest {file_path}: {str(e)}")
            continue
    
    return results
